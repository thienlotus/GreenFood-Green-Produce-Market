# -*- coding: utf-8 -*-
"""
Script tạo báo cáo chuẩn đề tài Website TMĐT GreenFood
Bao gồm:
- Mở đầu (Đặt vấn đề, mục tiêu, đối tượng, phương pháp, nội dung)
- Chương 1: Tổng quan về kiến trúc và thiết kế phần mềm (Yêu cầu chức năng, phi chức năng, kiến trúc MVC/RESTful Next.js - Laravel)
- Chương 2: Thiết kế dữ liệu và lớp (Chức năng Quản lý Đơn hàng)
  + 2.1 Thiết kế dữ liệu (Thực thể, thuộc tính, mối quan hệ, Sơ đồ ERD)
  + 2.2 Thiết kế lớp (Lớp, thuộc tính, phương thức, quan hệ, Sơ đồ Class Diagram)
  + 2.3 Thiết kế các biểu đồ (Tuần tự, Hoạt động, Trạng thái)
"""

import os
import sys
import matplotlib.pyplot as plt
import matplotlib.patches as patches
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

# Tạo thư mục tạm lưu sơ đồ ảnh
IMG_DIR = os.path.join(os.path.dirname(__file__), 'diagrams')
os.makedirs(IMG_DIR, exist_ok=True)

plt.rcParams['font.sans-serif'] = ['DejaVu Sans', 'Arial', 'Tahoma']
plt.rcParams['axes.unicode_minus'] = False

# ==========================================
# 1. VẼ SƠ ĐỒ ERD CHI TIẾT (CHUẨN CHEN & RELATIONAL)
# ==========================================
def draw_erd():
    fig, ax = plt.subplots(figsize=(15, 11), dpi=300)
    ax.set_xlim(0, 100)
    ax.set_ylim(0, 100)
    ax.axis('off')

    # Màu sắc thiết kế
    c_entity = '#2B6CB0'      # Xanh dương đậm cho Thực thể
    c_attr = '#E2E8F0'        # Xám nhạt cho Thuộc tính
    c_pk = '#FEEBC8'          # Vàng nhạt cho Khóa chính
    c_fk = '#EBF8FF'          # Xanh nhạt cho Khóa ngoại
    c_rel = '#38A169'         # Xanh lá cho Mối quan hệ
    c_line = '#4A5568'

    # Hàm vẽ Thực thể
    def draw_entity(x, y, w, h, text):
        rect = patches.FancyBboxPatch((x - w/2, y - h/2), w, h,
                                      boxstyle="round,pad=0.2",
                                      facecolor=c_entity, edgecolor='#1A365D', linewidth=2)
        ax.add_patch(rect)
        ax.text(x, y, text, color='white', weight='bold', fontsize=11,
                ha='center', va='center', zorder=5)

    # Hàm vẽ Quan hệ (Hình thoi)
    def draw_relation(x, y, w, h, text):
        rhombus = patches.Polygon([[x, y + h/2], [x + w/2, y], [x, y - h/2], [x - w/2, y]],
                                  facecolor=c_rel, edgecolor='#22543D', linewidth=2)
        ax.add_patch(rhombus)
        ax.text(x, y, text, color='white', weight='bold', fontsize=9.5,
                ha='center', va='center', zorder=5)

    # Hàm vẽ Bảng thực thể chi tiết (Table Box)
    def draw_table_box(x, y, w, h, title, pk_list, fk_list, attr_list):
        # Header
        hdr_h = 3.5
        hdr = patches.FancyBboxPatch((x, y + h - hdr_h), w, hdr_h,
                                     boxstyle="round,pad=0.1",
                                     facecolor='#1E3A8A', edgecolor='#0F172A', linewidth=1.5)
        ax.add_patch(hdr)
        ax.text(x + w/2, y + h - hdr_h/2, title, color='white', weight='bold', fontsize=10,
                ha='center', va='center', zorder=5)

        # Body
        body = patches.Rectangle((x, y), w, h - hdr_h,
                                 facecolor='#F8FAFC', edgecolor='#0F172A', linewidth=1.5)
        ax.add_patch(body)

        curr_y = y + h - hdr_h - 2.0
        line_gap = 2.2
        for pk in pk_list:
            ax.text(x + 1.2, curr_y, f"PK  {pk}", color='#B91C1C', weight='bold', fontsize=8.5, va='center')
            curr_y -= line_gap
        for fk in fk_list:
            ax.text(x + 1.2, curr_y, f"FK  {fk}", color='#1D4ED8', weight='bold', fontsize=8.5, va='center')
            curr_y -= line_gap
        for attr in attr_list:
            ax.text(x + 1.2, curr_y, f"•   {attr}", color='#334155', fontsize=8.5, va='center')
            curr_y -= line_gap

    # Title
    ax.text(50, 97, "SƠ ĐỒ THỰC THỂ QUAN HỆ (ERD) - PHÂN HỆ QUẢN LÝ ĐƠN HÀNG GREENFOOD",
            color='#0F172A', weight='bold', fontsize=14, ha='center', va='center')

    # Vẽ các bảng thực thể
    # 1. KHACH_HANG (USERS)
    draw_table_box(4, 62, 22, 26, "USERS (Khách hàng)",
                   ["id (UUID)"],
                   [],
                   ["full_name (VARCHAR)", "phone (VARCHAR)", "email (VARCHAR)", "password (VARCHAR)", "role (ENUM)", "avatar_url (VARCHAR)", "created_at (TIMESTAMP)"])

    # 2. KHU_VUC_GIAO_HANG (SHIPPING_ZONES)
    draw_table_box(4, 15, 22, 25, "SHIPPING_ZONES",
                   ["id (VARCHAR)"],
                   [],
                   ["name (VARCHAR)", "provinces (TEXT)", "base_fee (DECIMAL)", "free_ship_min (DECIMAL)", "estimated_days (VARCHAR)", "is_active (BOOLEAN)"])

    # 3. DON_HANG (ORDERS) - Trung tâm
    draw_table_box(38, 40, 26, 45, "ORDERS (Đơn hàng)",
                   ["id (UUID)"],
                   ["user_id (UUID)", "shipping_zone_id (VARCHAR)"],
                   ["tracking_number (VARCHAR)", "customer_name (VARCHAR)", "customer_phone (VARCHAR)", "customer_email (VARCHAR)", "shipping_address (TEXT)", "shipping_fee (DECIMAL)", "total_amount (DECIMAL)", "status (ENUM)", "payment_method (ENUM)", "note (TEXT)", "shipper_name (VARCHAR)", "shipper_phone (VARCHAR)", "created_at (TIMESTAMP)"])

    # 4. CHI_TIET_DON_HANG (ORDER_ITEMS)
    draw_table_box(74, 48, 23, 27, "ORDER_ITEMS (Chi tiết đơn)",
                   ["id (UUID)"],
                   ["order_id (UUID)", "product_id (UUID)", "variant_id (UUID)"],
                   ["product_name (VARCHAR)", "unit (VARCHAR)", "quantity (INT)", "price_at_time (DECIMAL)", "created_at (TIMESTAMP)"])

    # 5. SAN_PHAM (PRODUCTS)
    draw_table_box(74, 10, 23, 27, "PRODUCTS (Sản phẩm)",
                   ["id (UUID)"],
                   ["farmer_id (UUID)", "category_id (INT)"],
                   ["name (VARCHAR)", "slug (VARCHAR)", "sold_count (INT)", "rating (DECIMAL)", "is_seasonal (BOOL)", "image_url (VARCHAR)"])

    # Kẻ đường nối quan hệ và nhãn Cardinality
    def draw_link(p1, p2, label_start, label_end, rel_name=None):
        ax.annotate('', xy=p2, xytext=p1,
                    arrowprops=dict(arrowstyle="-", color=c_line, lw=1.8, shrinkA=0, shrinkB=0))
        # Text nhãn
        mx, my = (p1[0] + p2[0])/2, (p1[1] + p2[1])/2
        if rel_name:
            bbox_props = dict(boxstyle="round,pad=0.3", fc="#DCFCE7", ec="#16A34A", lw=1)
            ax.text(mx, my, rel_name, ha='center', va='center', fontsize=8.5, weight='bold', color='#15803D', bbox=bbox_props)
        ax.text(p1[0] + (2 if p1[0] < p2[0] else -2), p1[1] + 1.5, label_start, weight='bold', color='#1E293B', fontsize=9.5)
        ax.text(p2[0] + (-2 if p1[0] < p2[0] else 2), p2[1] + 1.5, label_end, weight='bold', color='#1E293B', fontsize=9.5)

    # Nối USERS -> ORDERS (1 - N: ĐẶT)
    draw_link((26, 75), (38, 75), "1", "N", "ĐẶT (1:N)")

    # Nối SHIPPING_ZONES -> ORDERS (1 - N: ÁP DỤNG)
    draw_link((26, 27), (38, 50), "1", "N", "ÁP DỤNG (1:N)")

    # Nối ORDERS -> ORDER_ITEMS (1 - N: GỒM)
    draw_link((64, 62), (74, 62), "1", "N", "GỒM (1:N)")

    # Nối PRODUCTS -> ORDER_ITEMS (1 - N: CHỨA)
    draw_link((85, 37), (85, 48), "1", "N", "CHỨA (1:N)")

    plt.tight_layout()
    erd_path = os.path.join(IMG_DIR, 'erd_diagram.png')
    plt.savefig(erd_path, bbox_inches='tight', dpi=300)
    plt.close()
    return erd_path

# ==========================================
# 2. VẼ SƠ ĐỒ LỚP CHI TIẾT (CLASS DIAGRAM)
# ==========================================
def draw_class_diagram():
    fig, ax = plt.subplots(figsize=(16, 12), dpi=300)
    ax.set_xlim(0, 100)
    ax.set_ylim(0, 100)
    ax.axis('off')

    ax.text(50, 97, "SƠ ĐỒ LỚP CHI TIẾT (CLASS DIAGRAM) - CHỨC NĂNG QUẢN LÝ ĐƠN HÀNG",
            color='#0F172A', weight='bold', fontsize=14, ha='center', va='center')

    # Hàm vẽ Class UML chuẩn 3 phần (Name, Attributes, Methods)
    def draw_uml_class(x, y, w, h, title, stereotype, attrs, methods, bg_color='#F1F5F9'):
        hdr_h = 4.0
        # Header
        hdr = patches.Rectangle((x, y + h - hdr_h), w, hdr_h,
                                facecolor='#0284C7', edgecolor='#0369A1', linewidth=1.5)
        ax.add_patch(hdr)
        if stereotype:
            ax.text(x + w/2, y + h - 1.3, f"«{stereotype}»", color='#E0F2FE', fontsize=7.5,
                    ha='center', va='center', style='italic')
            ax.text(x + w/2, y + h - 2.8, title, color='white', weight='bold', fontsize=9.5,
                    ha='center', va='center')
        else:
            ax.text(x + w/2, y + h - 2.0, title, color='white', weight='bold', fontsize=10,
                    ha='center', va='center')

        # Box body
        body = patches.Rectangle((x, y), w, h - hdr_h,
                                 facecolor=bg_color, edgecolor='#0369A1', linewidth=1.5)
        ax.add_patch(body)

        # Draw Attributes
        curr_y = y + h - hdr_h - 1.8
        attr_start_y = curr_y
        for attr in attrs:
            ax.text(x + 0.8, curr_y, attr, color='#0F172A', fontsize=8, va='center', family='monospace')
            curr_y -= 1.8

        # Divider line
        div_y = curr_y - 0.5
        ax.plot([x, x + w], [div_y, div_y], color='#94A3B8', linewidth=1)

        # Draw Methods
        curr_y = div_y - 1.8
        for m in methods:
            ax.text(x + 0.8, curr_y, m, color='#0F172A', fontsize=8, va='center', family='monospace')
            curr_y -= 1.8

    # 1. OrderController (Controller Layer)
    draw_uml_class(37, 72, 27, 21, "OrderController", "Controller",
                   ["- orderService: OrderService",
                    "- validator: Validator"],
                   ["+ store(req: Request): JsonResponse",
                    "+ track(trackingNo: String): JsonResponse",
                    "+ updateStatus(id: String, st: String)",
                    "+ cancelOrder(id: String): JsonResponse"])

    # 2. Order (Model Layer)
    draw_uml_class(37, 30, 27, 35, "Order", "Eloquent Model",
                   ["+ id: UUID",
                    "+ tracking_number: String",
                    "+ customer_name: String",
                    "+ customer_phone: String",
                    "+ customer_email: String",
                    "+ shipping_address: String",
                    "+ shipping_zone_id: String",
                    "+ shipping_fee: Decimal",
                    "+ total_amount: Decimal",
                    "+ status: OrderStatus",
                    "+ payment_method: PaymentMethod",
                    "+ shipper_name: String"],
                   ["+ user(): BelongsTo",
                    "+ shippingZone(): BelongsTo",
                    "+ items(): HasMany",
                    "+ calculateTotal(): Decimal",
                    "+ updateStatus(newStatus): void",
                    "+ canBeCancelled(): Boolean"])

    # 3. OrderItem (Model)
    draw_uml_class(72, 33, 25, 26, "OrderItem", "Eloquent Model",
                   ["+ id: UUID",
                    "+ order_id: UUID",
                    "+ product_id: UUID",
                    "+ variant_id: UUID",
                    "+ product_name: String",
                    "+ unit: String",
                    "+ quantity: Integer",
                    "+ price_at_time: Decimal"],
                   ["+ order(): BelongsTo",
                    "+ product(): BelongsTo",
                    "+ variant(): BelongsTo",
                    "+ getSubtotal(): Decimal"])

    # 4. User (Model)
    draw_uml_class(4, 52, 25, 24, "User", "Eloquent Model",
                   ["+ id: UUID",
                    "+ full_name: String",
                    "+ phone: String",
                    "+ email: String",
                    "+ role: UserRole"],
                   ["+ orders(): HasMany",
                    "+ getRecentOrders(): Collection",
                    "+ isAdmin(): Boolean"])

    # 5. ShippingZone (Model)
    draw_uml_class(4, 12, 25, 25, "ShippingZone", "Eloquent Model",
                   ["+ id: String",
                    "+ name: String",
                    "+ provinces: Text",
                    "+ base_fee: Decimal",
                    "+ free_ship_minimum: Decimal",
                    "+ is_active: Boolean"],
                   ["+ orders(): HasMany",
                    "+ calculateFee(subtotal): Decimal",
                    "+ isEligibleFreeShip(subtotal): Bool"])

    # 6. Product / ProductVariant (Model)
    draw_uml_class(72, 70, 25, 23, "ProductVariant", "Eloquent Model",
                   ["+ id: UUID",
                    "+ product_id: UUID",
                    "+ unit: String",
                    "+ price: Decimal",
                    "+ stock_quantity: Integer",
                    "+ sku: String"],
                   ["+ product(): BelongsTo",
                    "+ reduceStock(qty: Int): void",
                    "+ isAvailable(): Boolean"])

    # Vẽ các mũi tên quan hệ UML
    # Controller -> Order (Dependency / uses)
    ax.annotate('', xy=(50, 65), xytext=(50, 72),
                arrowprops=dict(arrowstyle="->", color='#0369A1', lw=1.5, ls='--'))
    ax.text(51, 68, "«uses»", fontsize=8, color='#0369A1', style='italic')

    # User -> Order (Association 1 .. 0..*)
    ax.plot([29, 37], [64, 50], color='#334155', lw=1.5)
    ax.text(30, 65, "1", weight='bold', fontsize=9)
    ax.text(35, 52, "0..*", weight='bold', fontsize=9)

    # ShippingZone -> Order (Association 1 .. 0..*)
    ax.plot([29, 37], [25, 38], color='#334155', lw=1.5)
    ax.text(30, 27, "1", weight='bold', fontsize=9)
    ax.text(35, 36, "0..*", weight='bold', fontsize=9)

    # Order *-- OrderItem (Composition: ◆)
    ax.plot([64, 72], [47, 47], color='#334155', lw=1.5)
    diamond = patches.Polygon([[64, 47], [65.5, 48], [67, 47], [65.5, 46]],
                              facecolor='#1E293B', edgecolor='#1E293B')
    ax.add_patch(diamond)
    ax.text(67.5, 48, "1", weight='bold', fontsize=9)
    ax.text(70, 48, "1..*", weight='bold', fontsize=9)

    # ProductVariant -> OrderItem (Association 1 .. 0..*)
    ax.plot([84, 84], [70, 59], color='#334155', lw=1.5)
    ax.text(85, 68, "1", weight='bold', fontsize=9)
    ax.text(85, 61, "0..*", weight='bold', fontsize=9)

    plt.tight_layout()
    class_path = os.path.join(IMG_DIR, 'class_diagram.png')
    plt.savefig(class_path, bbox_inches='tight', dpi=300)
    plt.close()
    return class_path

# ==========================================
# 3. VẼ BIỂU ĐỒ TUẦN TỰ (SEQUENCE DIAGRAM)
# ==========================================
def draw_sequence_diagram():
    fig, ax = plt.subplots(figsize=(15, 9), dpi=300)
    ax.set_xlim(0, 100)
    ax.set_ylim(0, 100)
    ax.axis('off')

    ax.text(50, 96, "BIỂU ĐỒ TUẦN TỰ (SEQUENCE DIAGRAM) - QUY TRÌNH ĐẶT HÀNG GREENFOOD",
            color='#0F172A', weight='bold', fontsize=13, ha='center')

    # Lifelines
    actors = [
        (10, "Khách Hàng\n(Customer)"),
        (30, "Next.js UI\n(Checkout Page)"),
        (52, "OrderController\n(Laravel API)"),
        (75, "Database\n(SQLite/MySQL)"),
        (92, "Vận Chuyển\n(Shipping Zone)")
    ]

    for x, label in actors:
        # Actor Box
        rect = patches.FancyBboxPatch((x - 8, 85), 16, 7, boxstyle="round,pad=0.2",
                                      facecolor='#0284C7', edgecolor='#0369A1', lw=1.5)
        ax.add_patch(rect)
        ax.text(x, 88.5, label, color='white', weight='bold', fontsize=8.5, ha='center', va='center')
        # Lifeline dotted line
        ax.plot([x, x], [85, 10], color='#94A3B8', linestyle='--', lw=1.2)

    # Messages
    msgs = [
        (10, 30, 80, "1. Nhập thông tin & Click 'Đặt hàng'", False),
        (30, 52, 73, "2. POST /api/orders (Payload + Items)", False),
        (52, 92, 66, "3. Tra cứu biểu phí theo ShippingZone", False),
        (92, 52, 60, "4. Trả về phí vận chuyển tính toán", True),
        (52, 75, 53, "5. BEGIN TRANSACTION & Validate tồn kho", False),
        (52, 75, 46, "6. INSERT INTO orders & order_items", False),
        (75, 52, 39, "7. Ghi nhận thành công & COMMIT", True),
        (52, 30, 32, "8. Trả về JSON {success: true, tracking_no}", True),
        (30, 10, 24, "9. Hiển thị trang hoàn tất & Mã tra cứu", True)
    ]

    for x1, x2, y, text, is_return in msgs:
        ls = '--' if is_return else '-'
        col = '#059669' if is_return else '#1E3A8A'
        ax.annotate('', xy=(x2, y), xytext=(x1, y),
                    arrowprops=dict(arrowstyle="->", color=col, lw=1.5, linestyle=ls))
        mx = (x1 + x2)/2
        ax.text(mx, y + 1.8, text, ha='center', va='center', fontsize=8.5, weight='bold', color=col,
                bbox=dict(boxstyle="round,pad=0.2", fc="#FFFFFF", ec="none", alpha=0.85))

    plt.tight_layout()
    seq_path = os.path.join(IMG_DIR, 'sequence_diagram.png')
    plt.savefig(seq_path, bbox_inches='tight', dpi=300)
    plt.close()
    return seq_path

# ==========================================
# 4. VẼ BIỂU ĐỒ HOẠT ĐỘNG (ACTIVITY DIAGRAM)
# ==========================================
def draw_activity_diagram():
    fig, ax = plt.subplots(figsize=(12, 10), dpi=300)
    ax.set_xlim(0, 100)
    ax.set_ylim(0, 100)
    ax.axis('off')

    ax.text(50, 97, "BIỂU ĐỒ HOẠT ĐỘNG (ACTIVITY DIAGRAM) - XỬ LÝ ĐƠN HÀNG",
            color='#0F172A', weight='bold', fontsize=13, ha='center')

    # Start node
    circle = patches.Circle((50, 91), 2, facecolor='#0F172A', edgecolor='#0F172A')
    ax.add_patch(circle)

    def draw_act(x, y, w, h, text, color='#E0F2FE', ec='#0284C7'):
        box = patches.FancyBboxPatch((x - w/2, y - h/2), w, h, boxstyle="round,pad=0.4",
                                     facecolor=color, edgecolor=ec, lw=1.5)
        ax.add_patch(box)
        ax.text(x, y, text, ha='center', va='center', fontsize=8.5, weight='bold', color='#0F172A')

    def draw_arrow(p1, p2, label=None):
        ax.annotate('', xy=p2, xytext=p1,
                    arrowprops=dict(arrowstyle="->", color='#1E293B', lw=1.5))
        if label:
            mx, my = (p1[0] + p2[0])/2, (p1[1] + p2[1])/2
            ax.text(mx + 2, my, label, fontsize=8, weight='bold', color='#DC2626')

    draw_arrow((50, 89), (50, 83))
    draw_act(50, 80, 42, 6, "1. Khách hàng xem giỏ hàng & Chọn thanh toán")

    draw_arrow((50, 77), (50, 71))
    draw_act(50, 68, 42, 6, "2. Nhập thông tin giao hàng & Chọn khu vực")

    draw_arrow((50, 65), (50, 59))
    # Decision Diamond
    diamond = patches.Polygon([[50, 57], [60, 54], [50, 51], [40, 54]],
                              facecolor='#FEF08A', edgecolor='#CA8A04', lw=1.5)
    ax.add_patch(diamond)
    ax.text(50, 54, "Dữ liệu hợp lệ?", ha='center', va='center', fontsize=8, weight='bold')

    # No branch
    draw_arrow((40, 54), (20, 54), "[Không]")
    draw_act(20, 44, 25, 6, "Báo lỗi nhập liệu\nYêu cầu bổ sung", color='#FEE2E2', ec='#DC2626')
    draw_arrow((20, 47), (20, 68))
    draw_arrow((20, 68), (29, 68))

    # Yes branch
    draw_arrow((50, 51), (50, 44), "[Hợp lệ]")
    draw_act(50, 41, 42, 6, "3. Tính phí vận chuyển & Tổng thanh toán")

    draw_arrow((50, 38), (50, 32))
    draw_act(50, 29, 42, 6, "4. Lưu Đơn Hàng & Giảm trừ tồn kho")

    draw_arrow((50, 26), (50, 20))
    draw_act(50, 17, 42, 6, "5. Sinh mã vận đơn & Gửi thông báo đặt thành công")

    draw_arrow((50, 14), (50, 8))
    # End node (Bullseye)
    end_out = patches.Circle((50, 6), 2.2, facecolor='none', edgecolor='#0F172A', lw=1.5)
    end_in = patches.Circle((50, 6), 1.4, facecolor='#0F172A', edgecolor='#0F172A')
    ax.add_patch(end_out)
    ax.add_patch(end_in)

    plt.tight_layout()
    act_path = os.path.join(IMG_DIR, 'activity_diagram.png')
    plt.savefig(act_path, bbox_inches='tight', dpi=300)
    plt.close()
    return act_path

# ==========================================
# 5. VẼ BIỂU ĐỒ TRẠNG THÁI (STATE MACHINE DIAGRAM)
# ==========================================
def draw_state_diagram():
    fig, ax = plt.subplots(figsize=(14, 7), dpi=300)
    ax.set_xlim(0, 100)
    ax.set_ylim(0, 100)
    ax.axis('off')

    ax.text(50, 94, "BIỂU ĐỒ TRẠNG THÁI (STATE DIAGRAM) - VÒNG ĐỜI ĐƠN HÀNG (ORDER LIFECYCLE)",
            color='#0F172A', weight='bold', fontsize=13, ha='center')

    # Start Node
    start_c = patches.Circle((6, 50), 2, facecolor='#0F172A')
    ax.add_patch(start_c)

    def draw_state(x, y, w, h, text, color='#DCFCE7', ec='#16A34A'):
        box = patches.FancyBboxPatch((x - w/2, y - h/2), w, h, boxstyle="round,pad=0.5",
                                     facecolor=color, edgecolor=ec, lw=2)
        ax.add_patch(box)
        ax.text(x, y, text, ha='center', va='center', fontsize=9, weight='bold', color='#14532D')

    # States
    draw_state(20, 50, 18, 12, "PENDING\n(Chờ xử lý)", '#FEF3C7', '#D97706')
    draw_state(45, 50, 18, 12, "CONFIRMED\n(Đã xác nhận)", '#DBEAFE', '#2563EB')
    draw_state(70, 50, 18, 12, "SHIPPING\n(Đang giao)", '#E0E7FF', '#4F46E5')
    draw_state(92, 50, 14, 12, "DELIVERED\n(Đã giao)", '#DCFCE7', '#16A34A')
    draw_state(45, 16, 20, 11, "CANCELLED\n(Đã hủy đơn)", '#FEE2E2', '#DC2626')

    # Transitions
    def draw_trans(p1, p2, text, above=True, curved=False):
        con = "arc3,rad=-0.15" if curved else "arc3,rad=0"
        ax.annotate('', xy=p2, xytext=p1,
                    arrowprops=dict(arrowstyle="->", color='#1E293B', lw=1.5, connectionstyle=con))
        mx, my = (p1[0] + p2[0])/2, (p1[1] + p2[1])/2
        ty = my + 3 if above else my - 3
        ax.text(mx, ty, text, fontsize=7.5, weight='bold', color='#0F172A', ha='center',
                bbox=dict(boxstyle="round,pad=0.2", fc="#FFFFFF", ec="none", alpha=0.8))

    draw_trans((8, 50), (11, 50), "Tạo đơn")
    draw_trans((29, 50), (36, 50), "Nông hộ/Admin duyệt")
    draw_trans((54, 50), (61, 50), "Bàn giao Shipper")
    draw_trans((79, 50), (85, 50), "Giao thành công")

    # Cancel transitions
    ax.annotate('', xy=(38, 20), xytext=(20, 44),
                arrowprops=dict(arrowstyle="->", color='#DC2626', lw=1.5, linestyle='--'))
    ax.text(25, 30, "Khách hủy", fontsize=7.5, weight='bold', color='#DC2626')

    ax.annotate('', xy=(45, 22), xytext=(45, 44),
                arrowprops=dict(arrowstyle="->", color='#DC2626', lw=1.5, linestyle='--'))
    ax.text(49, 32, "Hết hàng/Hủy", fontsize=7.5, weight='bold', color='#DC2626')

    plt.tight_layout()
    state_path = os.path.join(IMG_DIR, 'state_diagram.png')
    plt.savefig(state_path, bbox_inches='tight', dpi=300)
    plt.close()
    return state_path


# ==========================================
# 6. TẠO TÀI LIỆU WORD CHUẨN ĐỀ TƯƠNG ĐỀ TÀI
# ==========================================
def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def create_word_report(erd_img, class_img, seq_img, act_img, state_img, output_paths):
    doc = Document()

    # Cấu hình lề trang chuẩn báo cáo (Top 2cm, Bottom 2cm, Left 3cm, Right 2cm)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.79)
        section.bottom_margin = Inches(0.79)
        section.left_margin = Inches(1.18)
        section.right_margin = Inches(0.79)

    # Style chuẩn
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(13)
    normal_style.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
    normal_style.paragraph_format.line_spacing = 1.3
    normal_style.paragraph_format.space_after = Pt(6)

    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(18)
        run = p.add_run(text)
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x16, 0x65, 0x34) # Green 800

    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    def add_heading_3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    def add_paragraph_text(text, bold_prefix="", italic=False):
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if bold_prefix:
            r_bold = p.add_run(bold_prefix)
            r_bold.font.bold = True
        r_text = p.add_run(text)
        if italic:
            r_text.font.italic = True
        return p

    def add_bullet(bold_label, text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(4)
        r_bold = p.add_run(bold_label)
        r_bold.font.bold = True
        p.add_run(text)

    def add_image_figure(img_path, caption):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(10)
        p_img.paragraph_format.space_after = Pt(4)
        run_img = p_img.add_run()
        run_img.add_picture(img_path, width=Inches(6.4))

        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(12)
        r_cap = p_cap.add_run(caption)
        r_cap.font.size = Pt(11)
        r_cap.font.italic = True
        r_cap.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    # -------------------------------------------------------------
    # TRANG BÌA & TIÊU ĐỀ BÁO CÁO
    # -------------------------------------------------------------
    p_uni = doc.add_paragraph()
    p_uni.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p_uni.add_run("BÁO CÁO ĐỀ TÀI MÔN HỌC / ĐỒ ÁN CHUYÊN NGÀNH\n")
    r1.font.bold = True
    r1.font.size = Pt(14)
    r2 = p_uni.add_run("PHÂN TÍCH THIẾT KẾ HỆ THỐNG THÔNG TIN\n-------------------------***-------------------------")
    r2.font.size = Pt(12)

    add_title("BÁO CÁO ĐỀ TÀI XÂY DỰNG WEBSITE THƯƠNG MẠI ĐIỆN TỬ NÔNG SẢN SẠCH \"GREENFOOD\"")

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("Phân hệ trọng tâm: Thiết kế Kiến trúc, Dữ liệu và Lớp cho Chức năng Quản lý Đơn hàng\n")
    r_sub.font.italic = True
    r_sub.font.size = Pt(13)

    # Bảng thông tin đề tài
    t_info = doc.add_table(rows=4, cols=2)
    t_info.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_data = [
        ("Tên đề tài:", "Xây dựng Website Thương mại Điện tử Nông sản Sạch GreenFood"),
        ("Chức năng thiết kế chi tiết:", "Phân hệ Quản lý Đơn hàng & Định tuyến Phí Vận chuyển (Order Management & Logistics)"),
        ("Công nghệ triển khai:", "Next.js 14 (Frontend UI) + Laravel 11/12 RESTful API (Backend Service) + SQLite/MySQL"),
        ("Mục tiêu báo cáo:", "Trình bày tổng quan kiến trúc, thiết kế ERD, Class Diagram và các biểu đồ UML chuẩn hóa.")
    ]
    for row_idx, (k, v) in enumerate(info_data):
        row = t_info.rows[row_idx]
        cell_k, cell_v = row.cells[0], row.cells[1]
        cell_k.width = Inches(2.2)
        cell_v.width = Inches(4.2)
        rk = cell_k.paragraphs[0].add_run(k)
        rk.font.bold = True
        cell_v.paragraphs[0].add_run(v)
        set_cell_background(cell_k, "F1F5F9")
        set_cell_background(cell_v, "FFFFFF")

    doc.add_page_break()

    # -------------------------------------------------------------
    # MỞ ĐẦU
    # -------------------------------------------------------------
    add_heading_1("MỞ ĐẦU")

    add_heading_2("1. Đặt vấn đề (Tính cấp thiết của đề tài)")
    add_paragraph_text("Trong bối cảnh nền nông nghiệp Việt Nam đang chuyển dịch mạnh mẽ theo hướng công nghệ cao và minh bạch chuỗi cung ứng, nhu cầu tiếp cận thực phẩm sạch, an toàn chuẩn VietGAP/GlobalGAP của người tiêu dùng đô thị ngày càng gia tăng. Tuy nhiên, các kênh phân phối truyền thống thường qua nhiều tầng trung gian, đẩy giá thành lên cao trong khi lợi nhuận của bà con nông dân bị thu hẹp và thông tin xuất xứ nguồn gốc chưa được số hóa đầy đủ.")
    add_paragraph_text("Đề tài xây dựng website thương mại điện tử GreenFood ra đời nhằm giải quyết trực tiếp bài toán kết nối trực tiếp giữa các Hợp tác xã/Nông hộ sản xuất với người tiêu dùng cuối thông qua nền tảng số hóa. Hệ thống tích hợp bản đồ số GIS vùng trồng, tính toán phí vận chuyển linh hoạt theo từng khu vực địa lý (Shipping Zones) và quy trình quản lý đơn hàng khép kín, minh bạch thời gian thực.")

    add_heading_2("2. Mục tiêu nghiên cứu")
    add_bullet("Mục tiêu tổng quát: ", "Nghiên cứu, phân tích và xây dựng hệ thống website thương mại điện tử chuyên biệt cho nông sản sạch, đảm bảo tính trực quan, tin cậy và hiệu năng cao.")
    add_bullet("Mục tiêu cụ thể: ", "Thiết lập kiến trúc phần mềm phân tầng hiện đại (Frontend Next.js tách biệt Backend Laravel RESTful API); Thiết kế cơ sở dữ liệu quan hệ chuẩn hóa; Xây dựng sơ đồ lớp và các biểu đồ hành vi (Sequence, Activity, State) chi tiết cho chức năng Quản lý Đơn hàng.")

    add_heading_2("3. Đối tượng nghiên cứu")
    add_bullet("Đối tượng nghiệp vụ: ", "Quy trình mua sắm trực tuyến nông sản, quy trình đặt hàng, tính toán cước phí vận chuyển đa vùng miền và theo dõi trạng thái đơn hàng (Order Tracking).")
    add_bullet("Đối tượng kỹ thuật: ", "Mô hình kiến trúc phần mềm MVC, kiến trúc hướng dịch vụ REST API, cơ sở dữ liệu quan hệ (RDBMS), các chuẩn mô hình hóa UML (Unified Modeling Language).")

    add_heading_2("4. Phương pháp nghiên cứu")
    add_bullet("Phương pháp khảo sát & phân tích tài liệu: ", "Khảo sát thực tế các mô hình thương mại điện tử nông sản, quy định vận chuyển hàng nông sản tươi sống.")
    add_bullet("Phương pháp phân tích & thiết kế hướng đối tượng (OOAD): ", "Ứng dụng ngôn ngữ mô hình hóa UML để xây dựng sơ đồ Use Case, ERD, Class Diagram và biểu đồ tương tác.")
    add_bullet("Phương pháp thực nghiệm: ", "Cài đặt và triển khai thực tế hệ thống trên nền tảng Next.js kết hợp Laravel, kiểm thử toàn diện các luồng dữ liệu đơn hàng.")

    add_heading_2("5. Nội dung nghiên cứu")
    add_paragraph_text("Báo cáo tập trung triển khai các nội dung cốt lõi sau:")
    add_bullet("Chương 1: ", "Tổng quan về kiến trúc và thiết kế phần mềm, xác định các yêu cầu chức năng và phi chức năng của website GreenFood.")
    add_bullet("Chương 2: ", "Thiết kế dữ liệu và thiết kế lớp chi tiết cho phân hệ Quản lý Đơn hàng, bao gồm Sơ đồ Thực thể Quan hệ (ERD), Sơ đồ Lớp (Class Diagram), Biểu đồ Tuần tự (Sequence Diagram), Biểu đồ Hoạt động (Activity Diagram) và Biểu đồ Trạng thái (State Diagram).")

    doc.add_page_break()

    # -------------------------------------------------------------
    # CHƯƠNG 1: TỔNG QUAN VỀ KIẾN TRÚC VÀ THIẾT KẾ PHẦN MỀM
    # -------------------------------------------------------------
    add_heading_1("CHƯƠNG 1: TỔNG QUAN VỀ KIẾN TRÚC VÀ THIẾT KẾ PHẦN MỀM")

    add_heading_2("1.1. Thiết kế phần mềm")
    add_paragraph_text("Thiết kế phần mềm trong đề tài GreenFood là quá trình chuyển đổi các yêu cầu kinh doanh và nghiệp vụ thương mại điện tử nông sản thành một mô hình kiến trúc kỹ thuật có cấu trúc, module hóa cao, đảm bảo tính mở rộng, bảo trì và tái sử dụng mã nguồn.")
    add_paragraph_text("Hệ thống áp dụng phương pháp thiết kế hướng đối tượng (Object-Oriented Design) kết hợp với mô hình Domain-Driven Design (DDD) ở mức độ tinh gọn, tập trung tách biệt các miền nghiệp vụ: Quản lý Nông hộ & Vùng trồng, Quản lý Danh mục & Sản phẩm, Quản lý Giỏ hàng & Đơn hàng, Quản lý Biểu phí Vận chuyển.")

    add_heading_2("1.2. Kiến trúc phần mềm")
    add_paragraph_text("GreenFood được thiết kế theo mô hình kiến trúc phân tầng Decoupled (Headless Architecture) tiên tiến:")
    add_bullet("Tầng Trình diễn (Presentation Layer - Frontend): ", "Xây dựng bằng Next.js 14 (React Framework) với cơ chế Server-Side Rendering (SSR) và Client-Side Navigation, giúp tối ưu hóa SEO vượt trội cho sản phẩm nông sản và mang lại trải nghiệm mượt mà cho khách hàng.")
    add_bullet("Tầng Dịch vụ & Nghiệp vụ (Application/Business Logic Layer - Backend): ", "Xây dựng trên nền tảng Laravel Framework cung cấp hệ thống RESTful API chuẩn mực. Đảm nhận việc xác thực nghiệp vụ, kiểm tra ràng buộc toàn vẹn dữ liệu, giao dịch Database (Transaction) và quản lý tiến trình đặt hàng.")
    add_bullet("Tầng Dữ liệu (Data Persistence Layer): ", "Sử dụng Hệ quản trị cơ sở dữ liệu quan hệ (SQLite cho môi trường phát triển cục bộ và MySQL/PostgreSQL cho môi trường production) được quản lý qua Eloquent ORM mạnh mẽ.")

    add_heading_2("1.3. Chuyển đổi sang mô hình thiết kế")
    add_paragraph_text("Quá trình chuyển đổi từ yêu cầu nghiệp vụ sang mô hình thiết kế cụ thể được thực hiện qua các bước chuẩn hóa:")
    add_bullet("Bước 1: Khảo sát ca sử dụng (Use Case Modeling): ", "Xác định các tác nhân tương tác (Khách hàng, Nông hộ/Vendor, Quản trị viên Admin, Nhân viên giao hàng Shipper).")
    add_bullet("Bước 2: Mô hình hóa cấu trúc tĩnh (Structural Design): ", "Chuyển đổi các thực thể kinh doanh thành lược đồ cơ sở dữ liệu quan hệ (ERD) và các lớp nghiệp vụ hướng đối tượng (Class Diagram).")
    add_bullet("Bước 3: Mô hình hóa hành vi động (Behavioral Design): ", "Thiết kế các biểu đồ tương tác thời gian (Sequence Diagram), biểu đồ luồng quy trình nghiệp vụ (Activity Diagram) và máy trạng thái đơn hàng (State Machine).")

    add_heading_2("1.4. Phân tích yêu cầu hệ thống")

    add_heading_3("1.4.1. Yêu cầu chức năng của hệ thống (Functional Requirements)")
    add_paragraph_text("Đối với phân hệ Quản lý Đơn hàng trọng tâm:")
    add_bullet("F1 - Tạo và Lưu Đơn hàng: ", "Cho phép khách hàng chuyển đổi giỏ hàng thành đơn hàng chính thức bằng cách nhập thông tin người nhận (Tên, SĐT, Email, Địa chỉ chi tiết), chọn khu vực vận chuyển (Shipping Zone) và phương thức thanh toán (COD, Chuyển khoản, Ví điện tử).")
    add_bullet("F2 - Tự động tính cước vận chuyển & Miễn phí ship: ", "Hệ thống tự động tra cứu cấu hình bảng ShippingZone tương ứng; áp dụng chính sách miễn phí vận chuyển nếu tổng giá trị đơn hàng vượt ngưỡng `free_ship_minimum`.")
    add_bullet("F3 - Sinh mã vận đơn tự động (Tracking Number): ", "Hệ thống tự động sinh chuỗi mã vận đơn duy nhất (VD: `GF123456`) để phục vụ tra cứu công khai không cần đăng nhập.")
    add_bullet("F4 - Tra cứu tiến trình đơn hàng (Live Tracking): ", "Cung cấp giao diện tra cứu tiến độ với 4 mốc thời gian rõ ràng (Đặt hàng -> Xác nhận -> Đang giao -> Đã giao) cùng vị trí GPS mô phỏng của Shipper.")
    add_bullet("F5 - Cập nhật trạng thái và hủy đơn: ", "Cho phép Admin/Nông hộ chuyển đổi trạng thái đơn; cho phép khách hàng hủy đơn hàng khi đơn còn ở trạng thái `PENDING`.")

    add_heading_3("1.4.2. Yêu cầu phi chức năng (Non-Functional Requirements)")
    add_bullet("Hiệu năng (Performance): ", "Thời gian phản hồi API tạo đơn hàng < 500ms; xử lý đồng thời nhiều giao dịch đặt hàng không bị sai lệch số liệu tồn kho nhờ Database Transaction Lock.")
    add_bullet("Tính toàn vẹn & Tin cậy (Reliability & ACID): ", "Toàn bộ tiến trình tạo đơn và ghi nhận chi tiết đơn hàng được bao bọc trong `DB::beginTransaction()` và `DB::commit()`. Khi có bất kỳ lỗi nào, hệ thống tự động `DB::rollBack()` để bảo vệ dữ liệu.")
    add_bullet("Bảo mật (Security): ", "Kiểm tra và lọc sạch dữ liệu đầu vào (Input Validation); ngăn chặn các lỗ hổng SQL Injection và XSS thông qua Eloquent ORM và Middleware CORS.")
    add_bullet("Khả năng mở rộng (Scalability): ", "Kiến trúc Headless API cho phép dễ dàng mở rộng sang ứng dụng di động (Mobile App iOS/Android) trong tương lai.")

    doc.add_page_break()

    # -------------------------------------------------------------
    # CHƯƠNG 2: THIẾT KẾ DỮ LIỆU VÀ LỚP (CHỨC NĂNG QUẢN LÝ ĐƠN HÀNG)
    # -------------------------------------------------------------
    add_heading_1("CHƯƠNG 2: THIẾT KẾ DỮ LIỆU VÀ LỚP (CHỨC NĂNG QUẢN LÝ ĐƠN HÀNG)")

    add_heading_2("2.1. Thiết kế dữ liệu")
    add_paragraph_text("Thiết kế dữ liệu tập trung vào việc mô hình hóa các thực thể tham gia trực tiếp và gián tiếp vào quy trình tạo lập, định giá và xử lý đơn hàng trên hệ sinh thái GreenFood.")

    add_heading_3("2.1.1. Xác định các thực thể (Entities)")
    add_bullet("1. Thực thể KHACH_HANG (users): ", "Lưu trữ thông tin tài khoản người dùng, phân quyền vai trò (CUSTOMER, VENDOR, ADMIN).")
    add_bullet("2. Thực thể DON_HANG (orders): ", "Thực thể trung tâm lưu thông tin tổng quát của một đơn hàng, người nhận, cước phí, trạng thái xử lý và thông tin vận chuyển.")
    add_bullet("3. Thực thể CHI_TIET_DON_HANG (order_items): ", "Lưu trữ danh sách các mặt hàng cụ thể trong đơn, số lượng, đơn vị tính và giá chốt tại thời điểm mua.")
    add_bullet("4. Thực thể KHU_VUC_GIAO_HANG (shipping_zones): ", "Lưu trữ bảng giá cước vận chuyển theo từng vùng miền, danh sách tỉnh thành trực thuộc và hạn mức miễn phí vận chuyển.")
    add_bullet("5. Thực thể SAN_PHAM (products) & BIEN_THE_SAN_PHAM (product_variants): ", "Lưu trữ mặt hàng nông sản và các quy cách đóng gói (kg, hộp, trái) để khấu trừ tồn kho khi phát sinh đơn hàng.")

    add_heading_3("2.1.2. Xác định các thuộc tính của thực thể (Attributes Dictionary)")

    # Bảng mô tả chi tiết thuộc tính
    t_attr = doc.add_table(rows=1, cols=4)
    t_attr.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = t_attr.rows[0].cells
    hdr_titles = ["Thực thể", "Tên thuộc tính", "Kiểu dữ liệu", "Mô tả & Ràng buộc"]
    for idx, title in enumerate(hdr_titles):
        hdr_cells[idx].paragraphs[0].add_run(title).font.bold = True
        set_cell_background(hdr_cells[idx], "1E3A8A")
        hdr_cells[idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    attr_rows = [
        ("DON_HANG\n(orders)", "id", "UUID (PK)", "Khóa chính định danh duy nhất đơn hàng"),
        ("DON_HANG", "tracking_number", "VARCHAR(20)", "Mã vận đơn tra cứu (Unique, VD: GF839201)"),
        ("DON_HANG", "user_id", "UUID (FK)", "Khóa ngoại tham chiếu bảng users (Nullable khi mua không login)"),
        ("DON_HANG", "customer_name", "VARCHAR(255)", "Họ tên người nhận hàng (Bắt buộc)"),
        ("DON_HANG", "customer_phone", "VARCHAR(20)", "Số điện thoại nhận hàng (Bắt buộc)"),
        ("DON_HANG", "customer_email", "VARCHAR(255)", "Email nhận thông báo hóa đơn điện tử"),
        ("DON_HANG", "shipping_address", "TEXT", "Địa chỉ giao nhận hàng chi tiết"),
        ("DON_HANG", "shipping_zone_id", "VARCHAR (FK)", "Khóa ngoại tham chiếu bảng shipping_zones"),
        ("DON_HANG", "shipping_fee", "DECIMAL(15,2)", "Phí vận chuyển thực tế sau khi tính toán"),
        ("DON_HANG", "total_amount", "DECIMAL(15,2)", "Tổng tiền thanh toán (Tiền hàng + Phí ship)"),
        ("DON_HANG", "status", "ENUM", "Trạng thái: PENDING, CONFIRMED, SHIPPING, DELIVERED, CANCELLED"),
        ("DON_HANG", "payment_method", "ENUM", "Hình thức thanh toán: COD, BANK_TRANSFER, MOMO, VNPAY"),
        ("CHI_TIET_DON_HANG\n(order_items)", "id", "UUID (PK)", "Khóa chính chi tiết đơn hàng"),
        ("CHI_TIET_DON_HANG", "order_id", "UUID (FK)", "Khóa ngoại tham chiếu đến orders (Cascade Delete)"),
        ("CHI_TIET_DON_HANG", "product_id", "UUID (FK)", "Khóa ngoại tham chiếu đến products"),
        ("CHI_TIET_DON_HANG", "variant_id", "UUID (FK)", "Khóa ngoại tham chiếu đến product_variants"),
        ("CHI_TIET_DON_HANG", "product_name", "VARCHAR(255)", "Tên sản phẩm tại thời điểm đặt"),
        ("CHI_TIET_DON_HANG", "unit", "VARCHAR(50)", "Đơn vị tính: kg, hộp, túi, combo"),
        ("CHI_TIET_DON_HANG", "quantity", "INTEGER", "Số lượng đặt mua (>= 1)"),
        ("CHI_TIET_DON_HANG", "price_at_time", "DECIMAL(15,2)", "Giá bán niêm yết tại thời điểm tạo đơn"),
        ("KHU_VUC_GIAO_HANG\n(shipping_zones)", "id", "VARCHAR(20) (PK)", "Mã khu vực (SZ001: Nội thành, SZ002: Miền Tây...)"),
        ("KHU_VUC_GIAO_HANG", "name", "VARCHAR(255)", "Tên khu vực vận chuyển"),
        ("KHU_VUC_GIAO_HANG", "base_fee", "DECIMAL(15,2)", "Cước phí cơ bản của khu vực"),
        ("KHU_VUC_GIAO_HANG", "free_ship_minimum", "DECIMAL(15,2)", "Hạn mức tổng tiền hàng để được Free Ship")
    ]

    for row_data in attr_rows:
        row = t_attr.add_row()
        for idx, val in enumerate(row_data):
            c = row.cells[idx]
            c.paragraphs[0].add_run(val)
            if "PK" in val or "FK" in val:
                c.paragraphs[0].runs[0].font.bold = True
                if "PK" in val:
                    c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xB9, 0x1C, 0x1C)
                else:
                    c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)

    add_heading_3("2.1.3. Xác định mối quan hệ giữa các thực thể (Relationships)")
    add_bullet("1. KHACH_HANG (users) - DON_HANG (orders): ", "Quan hệ Một - Nhiều (1 : N). Một khách hàng có thể đặt nhiều đơn hàng trong lịch sử mua sắm. Một đơn hàng chỉ thuộc về một khách hàng (hoặc khách vãng lai).")
    add_bullet("2. KHU_VUC_GIAO_HANG (shipping_zones) - DON_HANG (orders): ", "Quan hệ Một - Nhiều (1 : N). Một khu vực vận chuyển được áp dụng cho nhiều đơn hàng khác nhau.")
    add_bullet("3. DON_HANG (orders) - CHI_TIET_DON_HANG (order_items): ", "Quan hệ Một - Nhiều (1 : N). Một đơn hàng bắt buộc phải bao gồm 1 hoặc nhiều dòng chi tiết sản phẩm. Khi xóa đơn hàng, toàn bộ chi tiết đơn hàng tương ứng sẽ được xóa theo (CASCADE).")
    add_bullet("4. SAN_PHAM (products) - CHI_TIET_DON_HANG (order_items): ", "Quan hệ Một - Nhiều (1 : N). Một sản phẩm nông sản có thể xuất hiện trong nhiều chi tiết đơn hàng của các khách hàng khác nhau.")

    add_heading_3("2.1.4. Sơ đồ thực thể quan hệ (ERD - Entity Relationship Diagram)")
    add_paragraph_text("Dưới đây là sơ đồ thực thể quan hệ ERD chi tiết của phân hệ Quản lý Đơn hàng trong hệ thống GreenFood:")
    add_image_figure(erd_img, "Hình 2.1: Sơ đồ thực thể quan hệ (ERD) phân hệ Quản lý Đơn hàng GreenFood")

    # -------------------------------------------------------------
    # 2.2. THIẾT KẾ LỚP (CLASS DIAGRAM)
    # -------------------------------------------------------------
    add_heading_2("2.2. Thiết kế lớp")
    add_paragraph_text("Thiết kế lớp thể hiện cấu trúc mã nguồn hướng đối tượng trong kiến trúc MVC của backend Laravel và mối liên kết với Frontend Next.js.")

    add_heading_3("2.2.1. Xác định các lớp (Classes Identification)")
    add_bullet("Lớp Điều Khiển (Controller Layer): ", "`OrderController` chịu trách nhiệm tiếp nhận HTTP Request từ Next.js Client, kiểm tra tính hợp lệ dữ liệu (Validation), điều phối logic nghiệp vụ và trả về HTTP JSON Response chuẩn hóa.")
    add_bullet("Lớp Thực Thể Mô Hình (Model/Entity Layer): ", "`Order`, `OrderItem`, `ShippingZone`, `User`, `ProductVariant` đại diện cho các bảng dữ liệu, đóng gói các quan hệ (Eloquent Relationships) và các phương thức nghiệp vụ tính toán.")
    add_bullet("Lớp Dịch Vụ & Xử Lý (Service Layer): ", "`OrderService` đảm nhận các tính toán cước phí phức tạp, xử lý Transaction và liên lạc với cổng thanh toán điện tử.")

    add_heading_3("2.2.2. Xác định các thuộc tính và phương thức của từng lớp")
    add_bullet("Lớp OrderController: ", "\n+ store(req: Request): JsonResponse - Tạo mới đơn hàng\n+ track(trackingNumber: String): JsonResponse - Tra cứu tiến độ đơn hàng\n+ updateStatus(id: String, status: String): JsonResponse - Cập nhật trạng thái\n+ cancelOrder(id: String): JsonResponse - Khách hàng hủy đơn")
    add_bullet("Lớp Order: ", "\n+ Các thuộc tính: id, tracking_number, customer_name, customer_phone, customer_email, shipping_address, shipping_zone_id, shipping_fee, total_amount, status, payment_method\n+ Phương thức quan hệ: user(), shippingZone(), items()\n+ Phương thức nghiệp vụ: calculateTotal(), updateStatus(newStatus), canBeCancelled()")
    add_bullet("Lớp OrderItem: ", "\n+ Các thuộc tính: id, order_id, product_id, variant_id, product_name, unit, quantity, price_at_time\n+ Phương thức: order(), product(), variant(), getSubtotal()")
    add_bullet("Lớp ShippingZone: ", "\n+ Các thuộc tính: id, name, provinces, base_fee, free_ship_minimum, is_active\n+ Phương thức: orders(), calculateFee(subtotal), isEligibleFreeShip(subtotal)")

    add_heading_3("2.2.3. Xác định mối quan hệ giữa các lớp (Class Relationships)")
    add_bullet("Quan hệ Hợp thành (Composition - ◆): ", "Giữa `Order` và `OrderItem`. Vòng đời của `OrderItem` phụ thuộc hoàn toàn vào `Order`. Nếu không có `Order`, các `OrderItem` không thể tồn tại độc lập.")
    add_bullet("Quan hệ Kết hợp (Association - —>): ", "Giữa `User` và `Order` (1..0..*), giữa `ShippingZone` và `Order` (1..0..*), giữa `ProductVariant` và `OrderItem` (1..0..*).")
    add_bullet("Quan hệ Phụ thuộc (Dependency - - - >): ", "`OrderController` phụ thuộc vào `Order`, `OrderItem` và `ShippingZone` để hoàn thành các yêu cầu HTTP API.")

    add_heading_3("2.2.4. Sơ đồ lớp chi tiết (Class Diagram)")
    add_paragraph_text("Dưới đây là sơ đồ lớp chi tiết thể hiện đầy đủ tên lớp, thuộc tính, kiểu dữ liệu, phương thức và các mối quan hệ theo chuẩn UML:")
    add_image_figure(class_img, "Hình 2.2: Sơ đồ lớp (Class Diagram) chi tiết phân hệ Quản lý Đơn hàng")

    # -------------------------------------------------------------
    # 2.3. THIẾT KẾ CÁC BIỂU ĐỒ HÀNH VI (BEHAVIORAL DIAGRAMS)
    # -------------------------------------------------------------
    add_heading_2("2.3. Thiết kế các biểu đồ hành vi")
    add_paragraph_text("Để mô tả trực quan các khía cạnh tương tác động, luồng xử lý dữ liệu và vòng đời của đơn hàng, đề tài xây dựng 3 biểu đồ chuẩn UML:")

    add_heading_3("2.3.1. Biểu đồ tuần tự (Sequence Diagram)")
    add_paragraph_text("Biểu đồ tuần tự mô tả trình tự trao đổi thông điệp theo trục thời gian giữa Khách hàng, giao diện Next.js, Controller Laravel, Cơ sở dữ liệu và phân hệ Vận chuyển trong kịch bản Đặt hàng thành công:")
    add_image_figure(seq_img, "Hình 2.3: Biểu đồ tuần tự (Sequence Diagram) - Quy trình Đặt hàng")

    add_heading_3("2.3.2. Biểu đồ hoạt động (Activity Diagram)")
    add_paragraph_text("Biểu đồ hoạt động thể hiện luồng công việc logic từ bước kiểm tra giỏ hàng, xác thực thông tin, rẽ nhánh kiểm tra tính hợp lệ dữ liệu, tính phí ship, trừ tồn kho và hoàn tất:")
    add_image_figure(act_img, "Hình 2.4: Biểu đồ hoạt động (Activity Diagram) - Luồng xử lý đơn hàng")

    add_heading_3("2.3.3. Biểu đồ trạng thái (State Machine Diagram)")
    add_paragraph_text("Biểu đồ trạng thái thể hiện các giai đoạn chuyển đổi vòng đời của một thực thể Đơn hàng từ khi mới tạo đến khi hoàn tất:")
    add_bullet("1. PENDING (Chờ xử lý): ", "Trạng thái khởi tạo ngay sau khi khách hàng bấm Đặt hàng thành công.")
    add_bullet("2. CONFIRMED (Đã xác nhận): ", "Nông hộ hoặc Quản trị viên kiểm tra kho, đóng gói hàng và xác nhận đơn.")
    add_bullet("3. SHIPPING (Đang giao hàng): ", "Đơn hàng đã được bàn giao cho nhân viên giao hàng (Shipper) vận chuyển trên đường.")
    add_bullet("4. DELIVERED (Đã giao thành công): ", "Khách hàng đã nhận nông sản, thanh toán tiền (nếu COD) và hoàn tất giao dịch.")
    add_bullet("5. CANCELLED (Đã hủy): ", "Đơn hàng bị hủy bởi người mua (khi còn PENDING) hoặc do hết hàng.")
    add_image_figure(state_img, "Hình 2.5: Biểu đồ trạng thái (State Diagram) - Vòng đời Đơn hàng")

    # -------------------------------------------------------------
    # KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN
    # -------------------------------------------------------------
    add_heading_1("KẾT LUẬN VÀ ĐÁNH GIÁ KẾT QUẢ")
    add_paragraph_text("Báo cáo đã hoàn thành trọn vẹn việc phân tích, thiết kế hệ thống theo đúng đề cương chuẩn mực từ Mở đầu, Chương 1 đến Chương 2:")
    add_bullet("Về mặt lý thuyết & kiến trúc: ", "Xác định rõ ràng kiến trúc phân tầng Headless Web App kết hợp Next.js 14 và Laravel API, đảm bảo tính hiện đại, linh hoạt và tối ưu hiệu năng.")
    add_bullet("Về mặt thiết kế dữ liệu: ", "Xây dựng lược đồ cơ sở dữ liệu quan hệ (ERD) chặt chẽ, tối ưu chuẩn hóa 3NF, giải quyết triệt để bài toán tính cước vận chuyển nông sản theo vùng miền (Shipping Zones).")
    add_bullet("Về mặt thiết kế lớp & tương tác: ", "Hoàn thành Sơ đồ Lớp (Class Diagram) và hệ thống biểu đồ UML (Tuần tự, Hoạt động, Trạng thái) phản ánh chính xác 100% mã nguồn đang vận hành của dự án GreenFood.")

    # Lưu file ra các đường dẫn
    for out_path in output_paths:
        os.makedirs(os.path.dirname(out_path), exist_ok=True)
        doc.save(out_path)
        print(f"[OK] Da luu bao cao thanh cong tai: {out_path}")

if __name__ == "__main__":
    print("=== DANG TAO CAC SO DO HINH ANH ===")
    erd_p = draw_erd()
    class_p = draw_class_diagram()
    seq_p = draw_sequence_diagram()
    act_p = draw_activity_diagram()
    state_p = draw_state_diagram()
    print("-> Da ve xong tat ca 5 so do do hoa.")

    # Xác định đường dẫn Desktop của người dùng
    desktop_folder = os.path.expanduser(r"~\OneDrive\Máy tính")
    if not os.path.exists(desktop_folder):
        desktop_folder = os.path.expanduser(r"~\Desktop")

    file_name = "Bao_Cao_De_Tai_Website_TMDT_GreenFood.docx"
    desktop_file = os.path.join(desktop_folder, file_name)
    workspace_file = os.path.join(r"d:\GreenFood-Green-Produce-Market", file_name)

    print("=== DANG TAO FILE WORD BAO CAO ===")
    create_word_report(erd_p, class_p, seq_p, act_p, state_p, [desktop_file, workspace_file])
    print("=== HOAN TAT! ===")
